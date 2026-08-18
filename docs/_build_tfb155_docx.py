#!/usr/bin/env python3
"""Збирає docs/TFB-155-gcp-migration.docx — покроковий опис міграції зі схемою.

Схему бере готовою з _build_tfb155_diagram.py, тож спершу треба запустити його.
Потребує python-docx (у системі його немає — ставиться в окремий venv):

    python3 -m venv /tmp/docxvenv && /tmp/docxvenv/bin/pip install python-docx Pillow
    /tmp/docxvenv/bin/python docs/_build_tfb155_diagram.py
    /tmp/docxvenv/bin/python docs/_build_tfb155_docx.py
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

INK = RGBColor(0x1A, 0x23, 0x7E)
MUTED = RGBColor(0x54, 0x6E, 0x7A)

DOCS = Path(__file__).parent
PNG = DOCS / "TFB-155-gcp-migration.png"

# (номер, заголовок, [абзаци], артефакт)
STEPS = [
    ("1", "Створити VM у межах Always Free", [
        "e2-micro в us-west1, us-central1 або us-east1 — це єдині регіони Always Free, — "
        "завантажувальний диск pd-standard на 30 GB, мережевий tier STANDARD.",
        "Пастка тут в дефолтах веб-консолі: вона пропонує pd-balanced, який у безкоштовний "
        "ліміт не входить, і машина починає тихо списувати кошти.",
        "Друга пастка — дефіцит самого заліза: чотири зони us-central1 поспіль відповіли "
        "ZONE_RESOURCE_POOL_EXHAUSTED, інстанс піднявся аж у us-east1-b. Це не помилка "
        "конфігурації, а привід перебирати зони — але лише в межах трьох безкоштовних регіонів.",
    ], "gcloud compute instances create"),
    ("2", "Підготувати систему під 1 GB", [
        "Swap-файл на 1 GB із vm.swappiness=10: він потрібен не для постійної роботи, а як "
        "подушка на піки — перезапуск JVM поверх ще живого старого контейнера, pg_dump, "
        "прогін Liquibase.",
        "Далі — Docker, каталог /opt/bookamore під git-checkout і .env з правами 600, "
        "заповнений із deploy/gcp/env.example.",
    ], "deploy/gcp/env.example"),
    ("3", "Винести збірку образів у CI", [
        "Головна зміна всієї міграції. Раніше сервер робив docker compose up --build, тобто "
        "збирав Maven і npm у себе. На 1 GB RAM це гарантований OOM.",
        "Тепер job build-images у GitHub Actions збирає обидва образи, публікує їх у GHCR "
        "тегами :latest і :sha, а VM виконує лише docker compose pull і up -d. "
        "Тег :sha дає відтворюваний деплой і тривіальний відкат.",
    ], ".github/workflows/deploy.yml"),
    ("4", "Ужати рантайм під ліміт памʼяті", [
        "JVM: -Xmx256m -Xss256k, SerialGC замість G1 (менше і памʼяті, і CPU на shared vCPU), "
        "TieredStopAtLevel=1, ExitOnOutOfMemoryError — щоб контейнер помирав і рестартував, "
        "а не зависав.",
        "Postgres: shared_buffers=48MB замість дефолтних 128MB, вимкнені паралельні воркери, "
        "max_connections=25. Пул Hikari обмежений пʼятьма зʼєднаннями.",
        "Дані — в іменованих томах bookamore_prod_db_data і bookamore_prod_uploads, тому "
        "переживають перестворення контейнерів. Логи json-file обмежені за розміром: 30 GB "
        "диска теж входять у безкоштовний ліміт.",
    ], "docker-compose.gcp.yaml"),
    ("5", "Полагодити X-Forwarded-Proto", [
        "Найтонше місце міграції. За тунелем nginx фронтенду слухає звичайний http, тож "
        "$scheme завжди дорівнює \"http\". Spring будував з нього redirect_uri вигляду "
        "http://www.bookamore.store/login/oauth2/code/google, а Google такий колбек "
        "відхиляє як невідповідний зареєстрованому.",
        "Рішення — директива map, яка бере схему з вхідного заголовка X-Forwarded-Proto "
        "(його виставляє cloudflared) і лише за його відсутності відкочується на $scheme. "
        "Перевірено локально: із заголовком redirect_uri стає https://, без нього — http://.",
        "Практичного ефекту зараз не видно — соціальний логін не працює ні в одному "
        "середовищі через заглушки замість креденшелів (див. «Межі рішення»). Але без цього "
        "фіксу він не запрацював би й після того, як застосунки в консолях зʼявляться.",
    ], "frontend/nginx.conf"),
    ("6", "Підняти Cloudflare Tunnel", [
        "cloudflared працює контейнером у тому ж compose-стеку і ходить до frontend:80 по "
        "внутрішній docker-мережі. Токен коннектора передається змінною TUNNEL_TOKEN, а не "
        "аргументом команди, інакше він світився б у docker ps і ps aux.",
        "Наслідок: на хості не публікується жоден порт. Тунель ініціюється зсередини VM, "
        "тому вхідні правила фаєрвола для 80/443 просто не потрібні.",
    ], "cloudflared у docker-compose.gcp.yaml"),
    ("7", "Завести власний домен у Cloudflare", [
        "Початковий план вів alt-web.biz.ua у Cloudflare, щоб bookamore.alt-web.biz.ua став "
        "CNAME на тунель. Від нього відмовились: та зона несе пошту (MX, SPF, DKIM-селектор "
        "dkim) і кілька сторонніх сайтів на 185.174.220.11, тож зміна NS ризикувала всім "
        "доменом заради одного піддомену.",
        "PROD натомість припаркований на власному bookamore.store — зоні, де до цього була "
        "сама лише парковка реєстратора, тому blast radius нульовий. Записи тунелю Cloudflare "
        "створює сам; парковочні A @ 162.255.119.47 і CNAME www на parkingpage треба зняти вручну.",
        "Канонічний хост — з www, апекс віддає 301 на нього через Dynamic Redirect Rule зі "
        "збереженням шляху. Два рівноправні хости не годяться: CLIENT_URL у беку — одне "
        "значення, і OAuth2SuccessHandler підставляє його в redirect цілим рядком.",
        "TLS далі тримає Cloudflare: Universal SSL покриває і апекс, і *.bookamore.store, "
        "certbot і host nginx на новій машині не потрібні взагалі.",
    ], "Cloudflare Dashboard"),
    ("8", "Перенести дані та закрити периметр", [
        "База переїжджає через pg_dump -Fc / pg_restore --clean, фото — розпакуванням "
        "архіву безпосередньо в named volume.",
        "Після перевірки: пересвідчитись, що правил default-allow-http і default-allow-https "
        "не лишилось (у мережі, створеній без галочок веб-консолі, їх може не бути взагалі), "
        "поставити budget alert на $1 і через добу звірити Billing → Reports.",
    ], "deploy/gcp/README.md, розділи 6 і 9"),
]

RAM = [
    ("Сервіс", "Ліміт", "Виміряно в спокої", "Очікуваний пік"),
    ("backend (JVM)", "460 MB", "232 MB", "380–420 MB"),
    ("db (Postgres 16)", "192 MB", "52 MB", "120–150 MB"),
    ("cloudflared", "48 MB", "—", "~30 MB"),
    ("frontend (nginx)", "32 MB", "12 MB", "~15 MB"),
    ("Разом", "732 MB", "296 MB", ""),
]

FILES = [
    ("Файл", "Що змінилось"),
    ("docker-compose.gcp.yaml", "Новий prod-стек: образи з GHCR, без портів на хості, "
                                "ліміти памʼяті, тюнінг Postgres, іменовані томи, cloudflared."),
    (".github/workflows/deploy.yml", "Розділено на три джоби: збірка образів у GHCR, деплой "
                                     "на GCP, окремо старий VPS для DEV."),
    ("frontend/nginx.conf", "X-Forwarded-Proto береться з вхідного заголовка; gzip на origin; "
                            "кеш /img/ на 30 днів."),
    ("backend/.../application-prod.yaml", "Логування в prod знижено до INFO/WARN замість "
                                          "DEBUG для Spring Security."),
    ("deploy/gcp/README.md", "Ранбук: створення VM, swap, тунель, зона bookamore.store, "
                             "перенесення даних, бекапи, контроль білінгу, доля старого PROD."),
    ("deploy/gcp/env.example", "Шаблон .env для GCP: образи, токен тунелю, канонічний "
                               "CLIENT_URL з www, порожні OAuth-креденшели з поясненням."),
    ("CLAUDE.md", "Опис деплою приведено до факту: main їде на GCP за www.bookamore.store, "
                  "dev — на старий VPS, зона alt-web.biz.ua лишається в старого реєстратора."),
]

CHECKS = [
    "docker stats — сумарний RSS тримається нижче ~800 MB, swap майже не зайнятий.",
    "У відповіді www.bookamore.store є server: cloudflare і cf-ray — трафік іде саме тунелем.",
    "Апекс віддає 301 на www-форму зі збереженим шляхом: /offers не перетворюється на головну.",
    "На /assets/ і /img/ приходить заголовок cf-cache-status: HIT — статику віддає edge, не origin.",
    "Прямий запит на публічний IP машини не відповідає.",
    "Фото і записи БД лишаються на місці після docker compose down і up.",
    "Billing → Reports через добу після запуску показує нуль по проєкту.",
]

LIMITS = [
    ("DEV не переїжджає", "Два JVM у 1 GB не вміщуються, тому на e2-micro живе лише PROD. "
                          "Домен bookamore-dev.alt-web.biz.ua лишається A-записом на старий VPS, "
                          "і збірка для нього так само відбувається на сервері. Другої "
                          "безкоштовної e2-micro не буде: ліміт Always Free рахується на "
                          "billing account, а не на проєкт."),
    ("Соціальний логін не працює ніде", "У всіх середовищах стоять заглушки: prod_placeholder "
                                        "на старому VPS і на GCP, change_me на DEV. Це не наслідок "
                                        "переїзду — застосунків у Google Cloud Console і Meta for "
                                        "Developers у проєкту просто немає, тож і redirect URI "
                                        "додавати нема куди. Перевіряється одним запитом: "
                                        "curl -sSI /oauth2/authorization/google і погляд на "
                                        "client_id у Location."),
    ("Старий PROD лишається як гарячий запас", "bookamore.alt-web.biz.ua і далі віддається зі "
                                               "старого VPS, але поза CI: main деплоїться тільки "
                                               "на GCP. Код там відстає з кожним мержем, а БД "
                                               "власна — перша ж реєстрація на старому домені "
                                               "створить запис, якого на GCP не буде. Перед "
                                               "вимкненням звірити totalElements на /api/v1/offers "
                                               "обох копій."),
    ("-Xss256k лишається агресивним", "Значення взяте з умови задачі й перевірене: застосунок "
                                      "стартує, Liquibase проходить, API відповідає. Якщо колись "
                                      "зʼявиться StackOverflowError — це перший параметр, який "
                                      "треба послабити до 512k."),
]


def style_doc(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Cm(1.8)
        section.left_margin = section.right_margin = Cm(2.0)


def add_table(doc, rows, widths):
    table = doc.add_table(rows=0, cols=len(rows[0]))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r, row in enumerate(rows):
        cells = table.add_row().cells
        for c, val in enumerate(row):
            cells[c].width = Cm(widths[c])
            p = cells[c].paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(9.5)
            if r == 0 or (rows is RAM and r == len(rows) - 1):
                run.bold = True
    return table


doc = Document()
style_doc(doc)

title = doc.add_heading("Міграція Bookamore на GCP e2-micro + Cloudflare Tunnel", level=0)
title.runs[0].font.color.rgb = INK

meta = doc.add_paragraph()
meta_run = meta.add_run("TFB-155 · PROD на GCP за www.bookamore.store · "
                        "DEV і старий PROD лишаються на VPS у зоні alt-web.biz.ua")
meta_run.font.size = Pt(9.5)
meta_run.font.color.rgb = MUTED

doc.add_paragraph(
    "Мета — тримати проєкт на безкоштовному хостингу за нуль доларів на місяць. "
    "Два ліміти визначають усі інші рішення: 1 GB оперативної памʼяті, через що збірка "
    "образів переїжджає в CI, і 1 GB вихідного трафіку на місяць, через що статику має "
    "кешувати Cloudflare, а не віддавати origin."
)

doc.add_heading("Схема потоку", level=1)
if PNG.exists():
    doc.add_picture(str(PNG), width=Cm(17.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run("Конвеєр доставки, шлях запиту в рантаймі та порядок виконання")
    cap_run.font.size = Pt(9)
    cap_run.font.color.rgb = MUTED
    cap_run.italic = True
else:
    doc.add_paragraph(f"[схему не знайдено: {PNG.name} — спершу запустіть _build_tfb155_diagram.py]")

doc.add_page_break()

doc.add_heading("Порядок виконання", level=1)
for num, head, paras, artifact in STEPS:
    h = doc.add_heading(f"{num}. {head}", level=2)
    h.runs[0].font.color.rgb = INK
    for text in paras:
        doc.add_paragraph(text)
    art = doc.add_paragraph()
    art_label = art.add_run("Артефакт: ")
    art_label.bold = True
    art_label.font.size = Pt(9.5)
    art_val = art.add_run(artifact)
    art_val.font.size = Pt(9.5)
    art_val.font.name = "Consolas"
    art_val.font.color.rgb = MUTED

doc.add_heading("Бюджет памʼяті", level=1)
doc.add_paragraph(
    "e2-micro — це 1024 MB, з яких ядру та docker-демону лишається близько 200 MB. "
    "Колонка «виміряно» — реальний docker stats після прогону цього ж compose-файлу "
    "з тими самими лімітами і JVM-прапорцями."
)
add_table(doc, RAM, [5.0, 3.2, 4.4, 4.4])

doc.add_heading("Що змінилось у репозиторії", level=1)
add_table(doc, FILES, [6.0, 11.0])

doc.add_heading("Перевірка результату", level=1)
for item in CHECKS:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("Межі рішення", level=1)
for head, text in LIMITS:
    p = doc.add_paragraph()
    run = p.add_run(f"{head}. ")
    run.bold = True
    p.add_run(text)

out = DOCS / "TFB-155-gcp-migration.docx"
doc.save(out)
print(f"DOCX: {out}")
